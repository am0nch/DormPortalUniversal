"""
Generates DormPortal Universal proposal document as .docx
Run: python3 generate_proposal.py
Output: DormPortal_Proposal.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Styles helpers ────────────────────────────────────────────────────────────
def set_font(run, bold=False, size=11, color=None, italic=False):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1, color=(0, 51, 102)):
    p = doc.add_heading(level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0] if p.runs else p.add_run(text)
    run.text = text
    run.font.color.rgb = RGBColor(*color)
    run.font.size = Pt({1: 16, 2: 13, 3: 11.5}.get(level, 11))
    run.bold = True
    return p

def body(text, space_before=0, space_after=6, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(1 + level * 0.6)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        # background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '003366')
        tcPr.append(shd)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # data rows
    for i, row in enumerate(rows):
        r = t.add_row()
        fill = 'EBF0F8' if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate(row):
            cell = r.cells[j]
            cell.text = str(val)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(10)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(width)
    return t


# ═══════════════════════════════════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ASIA-PACIFIC INTERNATIONAL UNIVERSITY")
r.bold = True; r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Student Administration Department")
r.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Dormitory Administration Office")
r.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("DORM PORTAL UNIVERSAL")
r.bold = True; r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("A Browser-Based Dormitory Administration System")
r.font.size = Pt(14); r.italic = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Technical Proposal and System Documentation")
r.font.size = Pt(12); r.bold = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Submitted by:\nRichmond Panganiban Ilao\nDorm Administration Staff\n")
r.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f"Date: {datetime.date.today().strftime('%B %d, %Y')}")
r.font.size = Pt(11)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════
heading("1. Executive Summary")
body(
    "DormPortal Universal is a fully custom, browser-based dormitory administration "
    "system developed specifically for APIU Elijah Hall. It consolidates every "
    "operational workflow of the Dormitory Administration — room management, student "
    "records, attendance, incidents, inventory, staff scheduling, maintenance requests, "
    "and financial reporting — into a single, unified platform that works entirely "
    "offline without any server infrastructure or internet connection."
)
body(
    "The system is currently live and operational, serving the Dean's office as the "
    "primary administrative tool. It is deployed as a Progressive Web App (PWA) on "
    "GitHub Pages and is accessible on both desktop and mobile devices. All data is "
    "stored locally in the browser (IndexedDB / localStorage), ensuring complete "
    "data privacy and zero dependency on third-party cloud services."
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. BACKGROUND & PROBLEM STATEMENT
# ═══════════════════════════════════════════════════════════════════════════════
heading("2. Background and Problem Statement")
body(
    "Prior to this system, Dormitory Administration at APIU relied on fragmented "
    "paper-based processes and disconnected spreadsheets for managing room assignments, "
    "student records, utility billing, key tracking, and staff scheduling. This "
    "approach created the following operational challenges:"
)
for item in [
    "No centralized record system — information was scattered across multiple Excel files "
    "and paper logs, making cross-referencing time-consuming.",
    "Manual computation errors in utility billing, clearance fee calculation, and room "
    "availability tracking.",
    "No audit trail for incidents, room inspections, or key issuances — critical for "
    "accountability and student dispute resolution.",
    "Limited mobility — administrative tasks required physical presence at the Dean's "
    "office computer.",
    "Data loss risk — no structured backup protocol for dormitory records.",
    "No mechanism for RA/Monitor staff to report incidents or conduct room inspections "
    "digitally.",
]:
    bullet(item)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. SYSTEM OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
heading("3. System Overview")
body(
    "DormPortal Universal is a 20-module web application built with vanilla JavaScript "
    "(ES6+), HTML5, and CSS3. It requires no installation, no internet connection after "
    "initial load, and no server infrastructure. The application runs directly in any "
    "modern browser and functions as a full Progressive Web App (PWA) — it can be "
    "installed on a device's home screen and used completely offline."
)

heading("3.1 Technical Architecture", level=2)
add_table(
    ["Component", "Technology", "Purpose"],
    [
        ["Frontend", "Vanilla JS ES6+, HTML5, CSS3", "All UI and business logic"],
        ["Data Storage", "IndexedDB + localStorage", "Persistent offline data store"],
        ["Photo Storage", "IndexedDB (Blob)", "Student and asset photos, no size limit"],
        ["Offline Support", "Service Worker (Cache API)", "Full offline functionality"],
        ["Backup", "JSON export / ZIP export / AES-256-GCM encrypted backup", "Data portability and recovery"],
        ["Deployment", "GitHub Pages (PWA)", "Zero-cost hosting, automatic HTTPS"],
        ["Security", "PBKDF2-SHA256 (100k iterations)", "Dean password hashing"],
    ],
    col_widths=[3.5, 5.5, 5.5]
)

doc.add_paragraph()

heading("3.2 System Access Levels", level=2)
add_table(
    ["Role", "Access Method", "Scope"],
    [
        ["Dean / Admin", "PBKDF2 password gate", "Full access to all 20 modules"],
        ["Resident Advisor (RA)", "PIN gate (RA Portal)", "Attendance, inspection, incidents, maintenance"],
        ["Monitor", "PIN gate (Monitor Portal)", "Incidents, key borrow log"],
    ],
    col_widths=[3.5, 5.0, 6.0]
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. MODULE DESCRIPTIONS
# ═══════════════════════════════════════════════════════════════════════════════
heading("4. Module Descriptions")
body("The system consists of 20 integrated modules covering all aspects of dormitory operations:")

modules = [
    ("Room Reservations", "Manages all bed slots across all rooms. Tracks check-in/check-out, room holds, clearance processing, away students, transfer-in records, waiting queue, and room availability states (available, held, soon-available, full, locked). Features 21 configurable data columns and full print support for Away and Queue lists."),
    ("Student Profiles", "Maintains a student directory with profile photos, contact details, and academic information. Supports A4 formatted profile card printing and bulk import via CSV/Excel."),
    ("Floor Plan", "Interactive visual room grid displaying real-time occupancy across all floors. Supports bathroom pairing configuration for utility billing."),
    ("Utilities", "Tracks electricity and hot water meter readings per room. Calculates per-student billing with shared cost splitting. Archives billing records by semester."),
    ("Reports", "13-tab report hub covering occupancy summaries, financial reports, attendance archives, clearance records, and more. All reports are print/PDF-ready."),
    ("Room Inspection", "Digitizes move-in and move-out checklists with photo capture. Automatically calculates damage charges and pre-fills clearance forms. Supports key issuance with deposit tracking."),
    ("Key Inventory", "Tracks key checkout, returns, and lost keys per room. Generates overdue alerts and maintains a history of all key transactions including semester-long assigned keys."),
    ("Inventory", "Asset management with Code 39 barcode generation. Supports asset models/templates, room-based asset assignment, and bedding inventory with semester filtering."),
    ("Attendance", "Nightly roll call module with curfew countdown timer. Imports approved leave requests from SARRA2. Tracks present, absent, and on-leave status per student."),
    ("Incidents", "Behavioral incident tracking with severity classification, follow-up date tracking, and SARRA2-compatible CSV export. Supports printable incident reports."),
    ("Student Admin", "Logs off-campus requests and Dean's assistance records. Tracks approval status and maintains a searchable history."),
    ("Dorm Workers", "HR register for Resident Advisors, Monitors, Janitors, and SWP workers. Tracks employment details, floor assignments, and coverage areas."),
    ("Staff Scheduling", "Weekly shift grid for all dorm staff categories. Auto-detects on-duty status based on current time. Generates printable A4 schedule sheets."),
    ("Plant Requests", "Maintenance request management with urgency classification (routine, urgent, emergency). Dashboard view with status tracking and resolution logging."),
    ("RA Portal", "Mobile-optimized launcher for Resident Advisors. PIN-gated access to attendance, room inspection, incident reporting, and maintenance requests. Designed for phone use during rounds."),
    ("Monitor Portal", "Mobile-optimized launcher for Monitor staff. PIN-gated access to incident reporting and key borrow log."),
    ("User Guide", "Interactive built-in documentation covering all 18 operational modules with step-by-step instructions, screenshots, and tips. Accessible offline."),
    ("Handbook", "Digital copy of the APIU Residence Hall Handbook including house rules, fine schedule, and policies. Searchable and print-ready."),
    ("Semester Registry", "Centralized semester configuration used across all modules. Manages active semester, semester history, and semester-based data filtering."),
    ("Backup & Restore", "Exports data as plain JSON, encrypted JSON (AES-256-GCM), or ZIP archive with photos. Supports full restoration from any backup format. Password hash is excluded from plain exports for security."),
]

for name, desc in modules:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5)
    r1 = p.add_run(f"{name}. ")
    r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.size = Pt(11)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. KEY FEATURES & CAPABILITIES
# ═══════════════════════════════════════════════════════════════════════════════
heading("5. Key Features and Capabilities")

heading("5.1 Offline-First Design", level=2)
body(
    "The application is designed to function without internet connectivity. Once loaded, "
    "all data operations (create, read, update, delete) occur locally in the browser. "
    "The Service Worker caches all application files, ensuring the system remains fully "
    "functional even during network outages — critical for a dormitory setting where "
    "reliable internet access cannot always be guaranteed."
)

heading("5.2 Data Security", level=2)
body("The system implements multiple security layers:")
for item in [
    "Dean password is hashed using PBKDF2-SHA256 with 100,000 iterations and a random 16-byte salt — the same standard used by password managers.",
    "Session authentication uses tab-scoped sessionStorage flags; the session clears automatically on browser close.",
    "5 failed login attempts trigger a 30-second lockout to prevent brute-force attacks.",
    "Encrypted backups use AES-256-GCM encryption keyed via PBKDF2 — unreadable without the password.",
    "Plain JSON and ZIP backups exclude the password hash, preventing offline brute-force attacks on exported data.",
]:
    bullet(item)

heading("5.3 Data Backup and Recovery", level=2)
body("Three backup formats are supported:")
add_table(
    ["Format", "Contains Photos", "Password Protected", "Recommended Use"],
    [
        ["Plain JSON", "No", "No", "Quick data backup, human-readable"],
        ["ZIP Archive", "Yes", "No", "Full backup including student/asset photos"],
        ["Encrypted JSON", "No", "Yes (AES-256-GCM)", "Secure off-device storage"],
    ],
    col_widths=[3.0, 3.5, 4.0, 4.0]
)

doc.add_paragraph()

heading("5.4 Cross-Tab Synchronization", level=2)
body(
    "When multiple browser tabs have the application open, data changes in one tab are "
    "automatically broadcast to all other open tabs via a custom IDB-based event system, "
    "ensuring all views remain consistent without a server."
)

heading("5.5 Print and Export", level=2)
body(
    "Every module that generates reports supports A4-optimized print layouts. Reports "
    "are generated client-side and sent directly to the browser's print dialog — no "
    "third-party PDF library required. The system supports:"
)
for item in [
    "Room clearance forms (A4 and A5 cost card variants)",
    "Weekly staff schedule sheets",
    "Nightly attendance sheets",
    "Student profile cards",
    "Incident reports",
    "Away and Queue lists with print buttons",
    "Key inventory reports",
    "Utility billing summaries",
]:
    bullet(item)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. DEVELOPMENT HISTORY & QUALITY
# ═══════════════════════════════════════════════════════════════════════════════
heading("6. Development History and Quality Assurance")

body(
    "DormPortal Universal has undergone continuous iterative development through multiple "
    "formal development sessions, each documented with commit records and session logs. "
    "The following table summarizes major milestones:"
)

add_table(
    ["Version / Phase", "Date", "Key Deliverables"],
    [
        ["v4.0 — Core System", "2026-06-17", "IndexedDB schema, Snipe-IT inventory, room inspection rewrite"],
        ["v4.1 — Auth Redesign", "2026-06-18", "Dean-only PBKDF2 auth, M365 bypass removal, encrypted backup"],
        ["v4.2 — Export/Import", "2026-06-18", "Export modal, encrypted backup, mobile header fixes"],
        ["v4.3 — ZIP Backup", "2026-06-19", "ZIP export/import with photos, password modal Change/Remove toggle"],
        ["v4.4 — Security Hardening", "2026-06-20", "42 XSS vulnerabilities resolved (CodeQL audit), idempotent archive migration"],
        ["v4.5 — Bug Audit", "2026-06-21", "6 critical bugs fixed including ZIP binary parsing, photo display, semester dropdown"],
        ["v4.6 — Auth & Privacy Fixes", "2026-06-21", "Semester header race fix, dormPwdHash stripped from plain exports"],
    ],
    col_widths=[4.5, 3.0, 7.0]
)

doc.add_paragraph()

body(
    "The codebase has been subjected to GitHub CodeQL static analysis scanning. All 42 "
    "identified XSS (Cross-Site Scripting) vulnerabilities have been resolved as of "
    "version 4.4. The repository currently has zero open code scanning alerts."
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. SYSTEM SCALE
# ═══════════════════════════════════════════════════════════════════════════════
heading("7. System Scale and Codebase")
body(
    "The following table shows the current size of each file in the system, reflecting "
    "the depth and completeness of each module:"
)

add_table(
    ["File", "Lines of Code", "Module"],
    [
        ["index.html", "2,134", "Main menu, auth, backup/restore, semester registry"],
        ["dorm-db.js", "984", "Central data layer (all IDB/localStorage operations)"],
        ["room-reservations.html", "2,221", "Room & bed management"],
        ["inventory.html", "2,717", "Asset tracking & bedding"],
        ["userguide.html", "3,035", "Interactive documentation"],
        ["handbook.html", "1,873", "Residence Hall Handbook"],
        ["reports.html", "1,764", "13-tab report hub"],
        ["room-inspection.html", "1,592", "Move-in/out checklists"],
        ["key-inventory.html", "1,584", "Key tracking"],
        ["student-profiles.html", "1,246", "Student directory"],
        ["staff-scheduling.html", "1,274", "Staff scheduling"],
        ["attendance.html", "1,234", "Nightly roll call"],
        ["student-admin.html", "1,012", "Off-campus & assistance logs"],
        ["incidents.html", "996", "Behavioral incident tracking"],
        ["dorm-workers.html", "757", "HR register"],
        ["plant-requests.html", "719", "Maintenance requests"],
        ["utilities.html", "688", "Utility billing"],
        ["ra-portal.html", "605", "RA mobile launcher"],
        ["monitor-portal.html", "571", "Monitor mobile launcher"],
        ["floor-plan.html", "514", "Visual room grid"],
        ["TOTAL", "~29,302", "—"],
    ],
    col_widths=[5.0, 3.5, 6.0]
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 8. BENEFITS TO THE DEPARTMENT
# ═══════════════════════════════════════════════════════════════════════════════
heading("8. Benefits to the Department")

benefits = [
    ("Zero Infrastructure Cost",
     "No server, no database license, no cloud subscription. The system runs entirely "
     "in the browser and is hosted for free on GitHub Pages with automatic HTTPS."),
    ("Complete Operational Integration",
     "All dormitory workflows — from room assignment to utility billing to incident "
     "reporting — are unified in one system. No more switching between spreadsheets."),
    ("Offline Reliability",
     "The system works without internet. Power outages or network disruptions do not "
     "affect administrative operations."),
    ("Mobile Accessibility",
     "RA and Monitor portals are designed for phone screens, enabling staff to submit "
     "reports from anywhere in the dormitory building."),
    ("Data Integrity",
     "Structured data model with validation rules eliminates common manual entry errors "
     "in room assignments, billing calculations, and student records."),
    ("Audit Trail",
     "All incidents, key transactions, room inspections, and clearance records are "
     "permanently logged with timestamps, supporting accountability and dispute resolution."),
    ("SARRA2 Integration",
     "Attendance leave imports and incident exports are compatible with the university's "
     "SARRA2 system, reducing double-entry work for the Dean's office."),
    ("Built-in Documentation",
     "The system ships with a comprehensive interactive User Guide and the full "
     "Residence Hall Handbook, accessible offline at any time."),
    ("Scalability",
     "The modular architecture allows new features to be added without affecting "
     "existing functionality. The system is designed to eventually integrate with SARRA2."),
]

for title, desc in benefits:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5)
    r1 = p.add_run(f"{title}. ")
    r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.size = Pt(11)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 9. PROPOSED NEXT STEPS
# ═══════════════════════════════════════════════════════════════════════════════
heading("9. Proposed Next Steps")
body("The following improvements are recommended for the next development phase:")

next_steps = [
    "Formal onboarding and training session for Dean and RA/Monitor staff",
    "M365/OneDrive integration for automated cloud backup of exported data",
    "SARRA2 bi-directional sync for student enrollment data",
    "Room reservation online request portal for students",
    "Dashboard analytics module for semester-over-semester occupancy trends",
    "Multi-dorm support (currently configured for Elijah Hall only)",
    "Audit log viewer module for tracking all data changes with timestamps",
]
for item in next_steps:
    bullet(item)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 10. CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════
heading("10. Conclusion")
body(
    "DormPortal Universal represents a significant contribution to the operational "
    "efficiency of APIU's Dormitory Administration. Built entirely from scratch to "
    "address the specific workflows of Elijah Hall, the system has eliminated the "
    "fragmentation of paper-based and spreadsheet-based administration, providing a "
    "secure, offline-capable, mobile-accessible platform that the Dean's office and "
    "dorm staff can rely on daily."
)
body(
    "With 20 integrated modules, over 29,000 lines of code, and a security posture "
    "verified by GitHub CodeQL scanning, DormPortal Universal is a production-ready "
    "system that directly supports the Student Administration Department's mission of "
    "providing safe, organized, and accountable dormitory operations for APIU students."
)
body(
    "This system is submitted as a formal contribution to the Dormitory Administration "
    "with the intent of continued development and integration with APIU's existing "
    "administrative systems."
)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Respectfully submitted,\n\nRichmond Panganiban Ilao\nDormitory Administration")
r.font.size = Pt(11)

# ── Save ──────────────────────────────────────────────────────────────────────
output = "/home/richmond/DormPortalUniversal/DormPortal_Proposal.docx"
doc.save(output)
print(f"Saved: {output}")
